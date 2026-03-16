from __future__ import annotations

import getpass
import os
import re
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openai import OpenAI


SYSTEM_PROMPT = """
You are a senior recruitment advisor writing for Bradford & Marsh Consulting.

Produce a client-ready recruitment audit in polished British English. The report must read like an experienced human consultant wrote it. Avoid AI-sounding phrasing, repetitive sentence patterns, inflated language, and generic filler.

Style requirements:
- Be direct, specific and commercially credible.
- Prefer plain, professional language over dramatic or robotic wording.
- Vary sentence length naturally.
- Do not repeat stock phrases such as "operational drag", "pain points", "market competitiveness", "hard-edged correction", or similar consultant clichés unless the evidence clearly supports them.
- Make each recommendation practical, proportionate and tied to a business outcome.
- Where the client is doing something well, acknowledge it without overpraise.

Audit these 12 areas:
1. Recruitment strategy and workforce planning
2. Performance metrics and funnel conversion
3. Employer brand and market perception
4. Job adverts and job specifications
5. Sourcing and advertising process
6. Application handling and screening
7. Interview process quality
8. Decision making and offer process
9. Onboarding and early retention
10. Staff turnover risks
11. Candidate experience
12. Process ownership and accountability

For each section provide exactly these headings:
Score: X/10
Current state
Key risks
Commercial impact
Immediate actions
Structural improvements

Then include exactly these final sections:
Executive overview
Top 5 strengths
Top 5 problems
30 day plan
60 day plan
90 day plan
Overall recruitment score
Final verdict

Formatting and content rules:
- Keep each section concise and useful.
- Ground comments in the data provided.
- Explain what each metric suggests instead of simply restating it.
- Avoid copying phrases between sections.
- Write recommendations that a managing director or operations lead could act on next week.
- Use bullet points only where they improve readability.
""".strip()

API_KEY_FILE = Path.home() / ".openai_api_key.txt"
DEFAULT_BENCHMARK_FILE = Path.cwd() / "uk_recruitment_benchmark_framework.xlsx"
DEFAULT_OUTPUT_DIR = Path.cwd() / "audit_outputs"


SECTION_ORDER = [
    "Recruitment strategy and workforce planning",
    "Performance metrics and funnel conversion",
    "Employer brand and market perception",
    "Job adverts and job specifications",
    "Sourcing and advertising process",
    "Application handling and screening",
    "Interview process quality",
    "Decision making and offer process",
    "Onboarding and early retention",
    "Staff turnover risks",
    "Candidate experience",
    "Process ownership and accountability",
]

FINAL_SECTION_ORDER = [
    "Executive overview",
    "Top 5 strengths",
    "Top 5 problems",
    "30 day plan",
    "60 day plan",
    "90 day plan",
    "Overall recruitment score",
    "Final verdict",
]

PRIMARY = RGBColor(24, 38, 71)
SECONDARY = RGBColor(79, 98, 122)
ACCENT = RGBColor(193, 154, 107)
TEXT = RGBColor(36, 36, 36)
MUTED = RGBColor(102, 112, 122)
LIGHT_BG = "F6F8FB"
BORDER = "D9DFE8"


@dataclass
class AuditInputs:
    client_name: str
    company_name: str
    sector: str
    employee_count: str
    annual_hires: str
    locations: str
    assessment_period: str
    notes: str
    metrics: pd.DataFrame
    benchmark: pd.DataFrame


def save_api_key(api_key: str) -> None:
    API_KEY_FILE.write_text(api_key.strip(), encoding="utf-8")
    try:
        os.chmod(API_KEY_FILE, 0o600)
    except OSError:
        pass


def load_api_key() -> str | None:
    if API_KEY_FILE.exists():
        value = API_KEY_FILE.read_text(encoding="utf-8").strip()
        return value or None
    return None


def get_api_key_interactive() -> str:
    existing = load_api_key()
    if existing:
        return existing

    print("\nOpenAI API key not found.")
    api_key = getpass.getpass("Paste your OpenAI API key: ").strip()
    if not api_key:
        raise ValueError("No API key supplied.")
    save_api_key(api_key)
    print("API key saved locally for future runs.\n")
    return api_key


def ask(prompt_text: str, default: str = "") -> str:
    suffix = f" [{default}]" if default else ""
    response = input(f"{prompt_text}{suffix}: ").strip()
    return response if response else default


def ask_multiline(prompt_text: str) -> str:
    print(f"{prompt_text} (press Enter twice to finish)")
    lines: list[str] = []
    while True:
        line = input()
        if line == "":
            if lines:
                break
            else:
                continue
        lines.append(line)
    return "\n".join(lines)


def file_safe_name(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    cleaned = re.sub(r"_+", "_", cleaned).strip("._")
    return cleaned or "recruitment_audit"


def load_benchmark_data(path: Path) -> pd.DataFrame:
    if not path.exists():
        raise FileNotFoundError(
            f"Benchmark workbook not found at: {path}\n"
            "Place uk_recruitment_benchmark_framework.xlsx in the project folder "
            "or update DEFAULT_BENCHMARK_FILE."
        )

    excel_file = pd.ExcelFile(path)
    frames: list[pd.DataFrame] = []

    for sheet in excel_file.sheet_names:
        df = pd.read_excel(path, sheet_name=sheet)
        df["Source Sheet"] = sheet
        frames.append(df)

    benchmark = pd.concat(frames, ignore_index=True)
    benchmark.columns = [str(c).strip() for c in benchmark.columns]
    return benchmark


def stage_one_context() -> dict[str, str]:
    print("=" * 72)
    print("Stage 1 of 3 | Client profile")
    print("=" * 72)

    client_name = ask("Client contact name")
    company_name = ask("Company name")
    sector = ask("Sector")
    employee_count = ask("Approximate employee count")
    annual_hires = ask("Approximate annual hires")
    locations = ask("Main hiring location(s)")
    assessment_period = ask("Assessment period", "Last 12 months")
    notes = ask_multiline("Context notes about hiring challenges, structure or priorities")

    return {
        "client_name": client_name,
        "company_name": company_name,
        "sector": sector,
        "employee_count": employee_count,
        "annual_hires": annual_hires,
        "locations": locations,
        "assessment_period": assessment_period,
        "notes": notes,
    }


def stage_two_metrics() -> pd.DataFrame:
    print("\n" + "=" * 72)
    print("Stage 2 of 3 | Recruitment metrics")
    print("=" * 72)
    print("Enter figures where known. Leave blank if unavailable.\n")

    metric_questions = [
        ("Applications per vacancy", ""),
        ("Qualified applicants per vacancy", ""),
        ("Screen to interview conversion (%)", ""),
        ("Interview to offer conversion (%)", ""),
        ("Offer acceptance rate (%)", ""),
        ("Average time to shortlist (days)", ""),
        ("Average time to hire (days)", ""),
        ("Agency usage (%)", ""),
        ("Cost per hire (£)", ""),
        ("90 day attrition (%)", ""),
        ("12 month attrition (%)", ""),
        ("Hiring manager satisfaction (out of 10)", ""),
        ("Candidate satisfaction (out of 10)", ""),
        ("Percentage of hires from referrals", ""),
        ("Percentage of hires from direct sourcing", ""),
    ]

    rows = []
    for label, default in metric_questions:
        value = ask(label, default)
        rows.append({"Metric": label, "Client Value": value})

    return pd.DataFrame(rows)


def preview_metrics(metrics_df: pd.DataFrame) -> None:
    filled = metrics_df["Client Value"].replace("", pd.NA).dropna()
    completion = int((len(filled) / len(metrics_df)) * 100) if len(metrics_df) else 0

    print("\nRecruitment metrics captured:")
    print(f"- Fields completed: {len(filled)}/{len(metrics_df)} ({completion}%)")
    for _, row in metrics_df.iterrows():
        value = row["Client Value"] if str(row["Client Value"]).strip() else "Not provided"
        print(f"  • {row['Metric']}: {value}")


def stage_three_confirmation(client_data: dict[str, str], metrics_df: pd.DataFrame) -> None:
    print("\n" + "=" * 72)
    print("Stage 3 of 3 | Review before generating report")
    print("=" * 72)
    print(f"Client: {client_data['client_name']}")
    print(f"Company: {client_data['company_name']}")
    print(f"Sector: {client_data['sector']}")
    print(f"Employees: {client_data['employee_count']}")
    print(f"Annual hires: {client_data['annual_hires']}")
    print(f"Locations: {client_data['locations']}")
    print(f"Assessment period: {client_data['assessment_period']}")
    print(f"Notes provided: {'Yes' if client_data['notes'].strip() else 'No'}")
    preview_metrics(metrics_df)
    print("")


def collect_inputs(benchmark_df: pd.DataFrame) -> AuditInputs:
    client_data = stage_one_context()
    metrics_df = stage_two_metrics()
    stage_three_confirmation(client_data, metrics_df)

    confirm = ask("Generate the report now? (y/n)", "y").lower()
    if confirm not in {"y", "yes"}:
        print("Report generation cancelled.")
        sys.exit(0)

    return AuditInputs(
        client_name=client_data["client_name"],
        company_name=client_data["company_name"],
        sector=client_data["sector"],
        employee_count=client_data["employee_count"],
        annual_hires=client_data["annual_hires"],
        locations=client_data["locations"],
        assessment_period=client_data["assessment_period"],
        notes=client_data["notes"],
        metrics=metrics_df,
        benchmark=benchmark_df,
    )


def benchmark_snapshot_text(benchmark_df: pd.DataFrame, sector: str) -> str:
    if benchmark_df.empty:
        return "No benchmark data available."

    working = benchmark_df.copy()
    working = working.fillna("")

    matches = pd.DataFrame()
    if "Sector" in working.columns and sector.strip():
        matches = working[working["Sector"].astype(str).str.contains(sector, case=False, na=False)]

    sample = matches if not matches.empty else working
    sample = sample.head(12)

    lines = ["Benchmark reference sample:"]
    for _, row in sample.iterrows():
        compact = []
        for col in sample.columns[:6]:
            val = str(row[col]).strip()
            if val:
                compact.append(f"{col}: {val}")
        if compact:
            lines.append("- " + " | ".join(compact))

    return "\n".join(lines) if len(lines) > 1 else "No benchmark sample available."


def build_user_prompt(inputs: AuditInputs) -> str:
    metrics_lines = []
    for _, row in inputs.metrics.iterrows():
        value = str(row["Client Value"]).strip() or "Not provided"
        metrics_lines.append(f"- {row['Metric']}: {value}")

    return f"""
Client details
- Client contact: {inputs.client_name}
- Company name: {inputs.company_name}
- Sector: {inputs.sector}
- Approximate employee count: {inputs.employee_count}
- Approximate annual hires: {inputs.annual_hires}
- Main hiring location(s): {inputs.locations}
- Assessment period: {inputs.assessment_period}

Client context notes
{inputs.notes or "No additional notes provided."}

Recruitment metrics supplied
{chr(10).join(metrics_lines)}

{benchmark_snapshot_text(inputs.benchmark, inputs.sector)}

Write the full recruitment audit using the required headings and sequence from the system prompt.
Keep it practical, commercially grounded and suitable for presentation to a client.
""".strip()


def generate_audit_text(client: OpenAI, inputs: AuditInputs) -> str:
    response = client.responses.create(
        model="gpt-4.1",
        input=[
            {"role": "system", "content": SYSTEM_PROMPT},
            {"role": "user", "content": build_user_prompt(inputs)},
        ],
        temperature=0.7,
    )
    return response.output_text.strip()


def split_audit_sections(text: str) -> dict[str, str]:
    headings = SECTION_ORDER + FINAL_SECTION_ORDER
    pattern = r"(?im)^(%s)\s*$" % "|".join(re.escape(h) for h in headings)
    matches = list(re.finditer(pattern, text))

    sections: dict[str, str] = {}
    if not matches:
        sections["Full report"] = text.strip()
        return sections

    for index, match in enumerate(matches):
        heading = match.group(1).strip()
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        body = text[start:end].strip()
        sections[heading] = body

    return sections


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    new_run.append(rpr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    new_run.append(text_element)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_document_defaults(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = TEXT

    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in document.styles:
            style = document.styles[style_name]
            style.font.name = "Aptos"

    sec = document.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)

    section_props = sec._sectPr
    cols = section_props.xpath("./w:cols")
    if cols:
        cols[0].set(qn("w:space"), "720")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 10.0, color: RGBColor = TEXT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_cover_page(document: Document, inputs: AuditInputs) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(50)
    p.paragraph_format.space_after = Pt(8)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Recruitment Audit")
    run.font.name = "Aptos"
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = PRIMARY

    p2 = document.add_paragraph()
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run(inputs.company_name)
    run2.font.name = "Aptos"
    run2.bold = True
    run2.font.size = Pt(16)
    run2.font.color.rgb = SECONDARY

    p3 = document.add_paragraph()
    p3.paragraph_format.space_after = Pt(14)
    run3 = p3.add_run(f"Prepared for {inputs.client_name}")
    run3.font.name = "Aptos"
    run3.font.size = Pt(11)
    run3.font.color.rgb = MUTED

    table = document.add_table(rows=4, cols=2)
    table.autofit = True
    summary_rows = [
        ("Sector", inputs.sector or "Not provided"),
        ("Employee count", inputs.employee_count or "Not provided"),
        ("Annual hires", inputs.annual_hires or "Not provided"),
        ("Assessment period", inputs.assessment_period or "Not provided"),
    ]
    for i, (label, value) in enumerate(summary_rows):
        set_cell_text(table.cell(i, 0), label, bold=True, size=10.2, color=PRIMARY)
        shade_cell(table.cell(i, 0), LIGHT_BG)
        set_cell_text(table.cell(i, 1), value, size=10.2)

    document.add_paragraph("")
    p4 = document.add_paragraph()
    p4.paragraph_format.space_after = Pt(0)
    run4 = p4.add_run("Bradford & Marsh Consulting")
    run4.font.name = "Aptos"
    run4.bold = True
    run4.font.size = Pt(11)
    run4.font.color.rgb = ACCENT

    document.add_page_break()


def add_section_title(document: Document, title: str, level: int = 1) -> None:
    heading = document.add_paragraph()
    heading.style = document.styles[f"Heading {min(level,3)}"]
    heading.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    heading.paragraph_format.space_after = Pt(5)

    run = heading.add_run(title)
    run.font.name = "Aptos"
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11.5)
    run.font.color.rgb = PRIMARY if level == 1 else SECONDARY


def clean_lines(text: str) -> list[str]:
    text = text.replace("\r", "").strip()
    if not text:
        return []

    lines = [ln.strip() for ln in text.split("\n")]
    collapsed: list[str] = []
    buffer = []

    for line in lines:
        if not line:
            if buffer:
                collapsed.append(" ".join(buffer).strip())
                buffer = []
            continue
        if re.match(r"^(Score:\s*\d+/10)$", line, re.I):
            if buffer:
                collapsed.append(" ".join(buffer).strip())
                buffer = []
            collapsed.append(line)
            continue
        if re.match(r"^[A-Za-z0-9][A-Za-z0-9 /\-&()]{1,40}$", line) and line.lower() in {
            "current state",
            "key risks",
            "commercial impact",
            "immediate actions",
            "structural improvements",
        }:
            if buffer:
                collapsed.append(" ".join(buffer).strip())
                buffer = []
            collapsed.append(line)
            continue
        if line.startswith("- ") or line.startswith("• "):
            if buffer:
                collapsed.append(" ".join(buffer).strip())
                buffer = []
            collapsed.append(line)
            continue
        buffer.append(line)

    if buffer:
        collapsed.append(" ".join(buffer).strip())

    return collapsed


def render_section_body(document: Document, text: str) -> None:
    lines = clean_lines(text)
    current_heading = None

    for line in lines:
        if re.match(r"^Score:\s*\d+/10$", line, re.I):
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line)
            run.bold = True
            run.font.name = "Aptos"
            run.font.size = Pt(10.5)
            run.font.color.rgb = ACCENT
            current_heading = None
            continue

        if line.lower() in {
            "current state",
            "key risks",
            "commercial impact",
            "immediate actions",
            "structural improvements",
        }:
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(line)
            run.bold = True
            run.font.name = "Aptos"
            run.font.size = Pt(10.5)
            run.font.color.rgb = SECONDARY
            current_heading = line.lower()
            continue

        if line.startswith("- ") or line.startswith("• "):
            bullet = document.add_paragraph(style="List Bullet")
            bullet.paragraph_format.space_before = Pt(0)
            bullet.paragraph_format.space_after = Pt(0)
            bullet.paragraph_format.left_indent = Inches(0.15)
            bullet.paragraph_format.first_line_indent = Inches(0)
            run = bullet.add_run(line[2:].strip())
            run.font.name = "Aptos"
            run.font.size = Pt(10.2)
            run.font.color.rgb = TEXT
            continue

        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.08
        run = p.add_run(line)
        run.font.name = "Aptos"
        run.font.size = Pt(10.2)
        run.font.color.rgb = TEXT


def extract_numeric(value: str) -> float | None:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None

    match = re.search(r"-?\d+(\.\d+)?", text.replace(",", ""))
    if not match:
        return None
    try:
        return float(match.group())
    except ValueError:
        return None


def metric_lookup(metrics_df: pd.DataFrame, label: str) -> float | None:
    series = metrics_df.loc[metrics_df["Metric"] == label, "Client Value"]
    if series.empty:
        return None
    return extract_numeric(str(series.iloc[0]))


def create_chart_pack(metrics_df: pd.DataFrame, output_dir: Path) -> list[tuple[Path, str]]:
    charts: list[tuple[Path, str]] = []

    funnel_labels = [
        "Applications per vacancy",
        "Qualified applicants per vacancy",
        "Screen to interview conversion (%)",
        "Interview to offer conversion (%)",
        "Offer acceptance rate (%)",
    ]
    funnel_values = [metric_lookup(metrics_df, label) for label in funnel_labels]

    if any(v is not None for v in funnel_values):
        x_labels = [
            "Applications\nper vacancy",
            "Qualified\napplicants",
            "Screen to\ninterview %",
            "Interview to\noffer %",
            "Offer\nacceptance %",
        ]
        y_values = [v if v is not None else 0 for v in funnel_values]

        plt.figure(figsize=(8.0, 4.6))
        bars = plt.bar(x_labels, y_values)
        plt.title("Recruitment Funnel and Conversion Measures")
        plt.ylabel("Reported value")
        plt.xlabel("Recruitment stage or conversion point")
        plt.tight_layout()

        for bar, value in zip(bars, y_values):
            plt.text(
                bar.get_x() + bar.get_width() / 2,
                bar.get_height(),
                f"{value:.0f}",
                ha="center",
                va="bottom",
                fontsize=9,
            )

        path = output_dir / "chart_funnel_metrics.png"
        plt.savefig(path, dpi=220, bbox_inches="tight")
        plt.close()
        charts.append(
            (
                path,
                "Figure 1. Recruitment funnel and conversion measures. This chart shows the volume and conversion points across the hiring process, using the figures supplied by the client.",
            )
        )

    speed_labels = [
        "Average time to shortlist (days)",
        "Average time to hire (days)",
    ]
    speed_values = [metric_lookup(metrics_df, label) for label in speed_labels]
    if any(v is not None for v in speed_values):
        x_labels = ["Time to shortlist", "Time to hire"]
        y_values = [v if v is not None else 0 for v in speed_values]

        plt.figure(figsize=(7.6, 4.4))
        bars = plt.bar(x_labels, y_values)
        plt.title("Recruitment Process Speed")
        plt.ylabel("Days")
        plt.xlabel("Process stage")
        plt.tight_layout()

        for bar, value in zip(bars, y_values):
            plt.text(
                bar.get_x() + bar.get_width() / 2,
                bar.get_height(),
                f"{value:.0f} days",
                ha="center",
                va="bottom",
                fontsize=9,
            )

        path = output_dir / "chart_process_speed.png"
        plt.savefig(path, dpi=220, bbox_inches="tight")
        plt.close()
        charts.append(
            (
                path,
                "Figure 2. Recruitment process speed. This chart measures how long it takes to move from application to shortlist and from vacancy to accepted hire.",
            )
        )

    quality_labels = [
        "90 day attrition (%)",
        "12 month attrition (%)",
        "Hiring manager satisfaction (out of 10)",
        "Candidate satisfaction (out of 10)",
    ]
    quality_values = [metric_lookup(metrics_df, label) for label in quality_labels]
    if any(v is not None for v in quality_values):
        x_labels = [
            "90 day\nattrition %",
            "12 month\nattrition %",
            "Hiring manager\nsatisfaction",
            "Candidate\nsatisfaction",
        ]
        y_values = [v if v is not None else 0 for v in quality_values]

        plt.figure(figsize=(8.4, 4.6))
        bars = plt.bar(x_labels, y_values)
        plt.title("Retention and Stakeholder Satisfaction")
        plt.ylabel("Reported value")
        plt.xlabel("Retention or satisfaction measure")
        plt.tight_layout()

        for bar, value in zip(bars, y_values):
            suffix = "/10" if value <= 10 and bar.get_x() > 1.5 else "%"
            if "satisfaction" in x_labels[list(bars).index(bar)].lower():
                suffix = "/10"
            elif "attrition" in x_labels[list(bars).index(bar)].lower():
                suffix = "%"
            plt.text(
                bar.get_x() + bar.get_width() / 2,
                bar.get_height(),
                f"{value:.0f}{suffix}",
                ha="center",
                va="bottom",
                fontsize=9,
            )

        path = output_dir / "chart_retention_satisfaction.png"
        plt.savefig(path, dpi=220, bbox_inches="tight")
        plt.close()
        charts.append(
            (
                path,
                "Figure 3. Retention and stakeholder satisfaction. This chart shows early attrition alongside hiring manager and candidate feedback to indicate recruitment quality after offer and onboarding.",
            )
        )

    source_labels = [
        "Agency usage (%)",
        "Percentage of hires from referrals",
        "Percentage of hires from direct sourcing",
    ]
    source_values = [metric_lookup(metrics_df, label) for label in source_labels]
    if any(v is not None for v in source_values):
        x_labels = ["Agency usage %", "Referral hires %", "Direct sourcing %"]
        y_values = [v if v is not None else 0 for v in source_values]

        plt.figure(figsize=(7.8, 4.4))
        bars = plt.bar(x_labels, y_values)
        plt.title("Hiring Channel Mix")
        plt.ylabel("Percentage of hires or reliance")
        plt.xlabel("Sourcing channel")
        plt.tight_layout()

        for bar, value in zip(bars, y_values):
            plt.text(
                bar.get_x() + bar.get_width() / 2,
                bar.get_height(),
                f"{value:.0f}%",
                ha="center",
                va="bottom",
                fontsize=9,
            )

        path = output_dir / "chart_channel_mix.png"
        plt.savefig(path, dpi=220, bbox_inches="tight")
        plt.close()
        charts.append(
            (
                path,
                "Figure 4. Hiring channel mix. This chart shows how dependent the business is on agencies compared with referrals and direct sourcing activity.",
            )
        )

    return charts


def add_metrics_snapshot(document: Document, metrics_df: pd.DataFrame) -> None:
    add_section_title(document, "Client metrics snapshot", level=1)

    available = metrics_df.copy()
    available["Display Value"] = available["Client Value"].apply(
        lambda v: str(v).strip() if str(v).strip() else "Not provided"
    )

    table = document.add_table(rows=1, cols=2)
    table.autofit = True

    header_cells = table.rows[0].cells
    set_cell_text(header_cells[0], "Metric", bold=True, size=10.2, color=PRIMARY)
    set_cell_text(header_cells[1], "Client value", bold=True, size=10.2, color=PRIMARY)
    shade_cell(header_cells[0], LIGHT_BG)
    shade_cell(header_cells[1], LIGHT_BG)

    for _, row in available.iterrows():
        cells = table.add_row().cells
        set_cell_text(cells[0], str(row["Metric"]), size=10.0)
        set_cell_text(cells[1], str(row["Display Value"]), size=10.0)

    document.add_paragraph("")


def add_chart_section(document: Document, chart_paths: Iterable[tuple[Path, str]]) -> None:
    chart_paths = list(chart_paths)
    if not chart_paths:
        return

    add_section_title(document, "Charts and visual analysis", level=1)

    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(5)
    run = intro.add_run(
        "The following charts are based on the data supplied by the client and are labelled to show exactly what each figure is measuring."
    )
    run.font.name = "Aptos"
    run.font.size = Pt(10.2)
    run.font.color.rgb = TEXT

    for chart_path, caption in chart_paths:
        document.add_picture(str(chart_path), width=Inches(6.55))
        cap = document.add_paragraph()
        cap.paragraph_format.space_before = Pt(2)
        cap.paragraph_format.space_after = Pt(6)
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = cap.add_run(caption)
        run.italic = True
        run.font.name = "Aptos"
        run.font.size = Pt(9.2)
        run.font.color.rgb = MUTED


def build_word_report(inputs: AuditInputs, audit_text: str, output_path: Path, chart_paths: list[tuple[Path, str]]) -> None:
    sections = split_audit_sections(audit_text)

    document = Document()
    set_document_defaults(document)
    add_cover_page(document, inputs)
    add_metrics_snapshot(document, inputs.metrics)
    add_chart_section(document, chart_paths)

    add_section_title(document, "Audit findings", level=1)

    for heading in SECTION_ORDER:
        body = sections.get(heading, "No content generated for this section.")
        add_section_title(document, heading, level=2)
        render_section_body(document, body)

    add_section_title(document, "Recommendations and conclusion", level=1)

    for heading in FINAL_SECTION_ORDER:
        body = sections.get(heading, "No content generated for this section.")
        add_section_title(document, heading, level=2)
        render_section_body(document, body)

    document.add_section(WD_SECTION.NEW_PAGE)
    add_section_title(document, "Need Professional Help in Developing Your Architecture?", level=1)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Please contact me at ")
    run.font.name = "Aptos"
    run.font.size = Pt(10.5)
    run.font.color.rgb = TEXT
    add_hyperlink(p, "sammuti.com", "https://sammuti.com")
    run2 = p.add_run(" :)")
    run2.font.name = "Aptos"
    run2.font.size = Pt(10.5)
    run2.font.color.rgb = TEXT

    document.save(output_path)


def save_text_report(audit_text: str, output_path: Path) -> None:
    output_path.write_text(audit_text, encoding="utf-8")


def main() -> None:
    print("\nBradford & Marsh Consulting | Recruitment Audit Generator")
    print("-" * 72)

    benchmark_df = load_benchmark_data(DEFAULT_BENCHMARK_FILE)
    inputs = collect_inputs(benchmark_df)

    output_dir = DEFAULT_OUTPUT_DIR / file_safe_name(inputs.company_name or inputs.client_name)
    output_dir.mkdir(parents=True, exist_ok=True)

    api_key = get_api_key_interactive()
    client = OpenAI(api_key=api_key)

    print("\nGenerating narrative report...")
    audit_text = generate_audit_text(client, inputs)

    print("Building charts...")
    chart_paths = create_chart_pack(inputs.metrics, output_dir)

    base_name = file_safe_name(f"{inputs.company_name}_recruitment_audit")
    txt_path = output_dir / f"{base_name}.txt"
    docx_path = output_dir / f"{base_name}.docx"

    save_text_report(audit_text, txt_path)
    build_word_report(inputs, audit_text, docx_path, chart_paths)

    print("\nComplete.")
    print(f"Text report:  {txt_path}")
    print(f"Word report:  {docx_path}")
    if chart_paths:
        print("Charts:")
        for path, _ in chart_paths:
            print(f" - {path}")


if __name__ == "__main__":
    main()
