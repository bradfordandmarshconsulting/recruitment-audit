from __future__ import annotations

import json
import os
import re
import textwrap
from datetime import datetime
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
import pandas as pd
from matplotlib.ticker import MaxNLocator
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
BENCHMARK_FILE = BASE_DIR / "uk_recruitment_benchmark_framework.xlsx"
BENCHMARK_CSV_FILE = BASE_DIR / "uk_recruitment_benchmarks.csv"
OUTPUT_DIR = Path(os.environ.get("AUDIT_OUTPUT_DIR", "/tmp/BradfordMarshAI"))
BENCHMARK_ENV_VAR = "RECRUITMENT_BENCHMARK_FILE"

SECTION_ORDER = [
    "Recruitment strategy and workforce planning",
    "Performance metrics and funnel conversion",
    "Employer brand and market perception",
    "Job adverts and job specifications",
    "Sourcing and advertising process",
    "Application handling and screening",
    "Interview process quality",
    "Decision making and offer process",
    "Onboarding and early retention",
    "Staff turnover risks",
    "Candidate experience",
    "Process ownership and accountability",
]

FINAL_SECTION_ORDER = [
    "Executive overview",
    "Top 5 strengths",
    "Top 5 problems",
    "30 day plan",
    "60 day plan",
    "90 day plan",
    "Overall recruitment score",
    "Final verdict",
]

SECTION_IDS = [
    "recruitment_strategy_and_workforce_planning",
    "performance_metrics_and_funnel_conversion",
    "employer_brand_and_market_perception",
    "job_adverts_and_job_specifications",
    "sourcing_and_advertising_process",
    "application_handling_and_screening",
    "interview_process_quality",
    "decision_making_and_offer_process",
    "onboarding_and_early_retention",
    "staff_turnover_risks",
    "candidate_experience",
    "process_ownership_and_accountability",
]

SECTION_ID_TO_TITLE = dict(zip(SECTION_IDS, SECTION_ORDER))

SECTION_KEYS = {
    "current_state": "Current state",
    "key_risks": "Key risks",
    "commercial_impact": "Commercial impact",
    "immediate_actions": "Immediate actions",
    "structural_improvements": "Structural improvements",
}

BRAND_NAME = "Bradford & Marsh Consulting"
BRAND_STRAPLINE = "Recruitment advisory and operating model diagnostics"
REPORT_TITLE = "Recruitment Operating Model Audit"
CONFIDENTIAL_LABEL = "Confidential client report"
LETTERHEAD_TITLE = "BRADFORD & MARSH"
LETTERHEAD_SUBTITLE = "CONSULTING"
LETTERHEAD_COMPANY_NO = "Company Registration No. 17059554"
LETTERHEAD_REGISTERED_ADDRESS = (
    "Registered Company Address: Riverside Mill, Mountbatten Way, Congleton, Cheshire, CW12 1DY"
)
LETTERHEAD_PHONE = "T. 01260 544934"
LETTERHEAD_WEBSITE = "www.bradfordandmarsh.co.uk"

PRIMARY_HEX = "182647"
SECONDARY_HEX = "4F627A"
ACCENT_HEX = "C19A6B"
TEXT_HEX = "000000"
MUTED_HEX = "66707A"
WHITE_HEX = "FFFFFF"
LIGHT_BG = "F6F8FB"
SOFT_BG = "F8FAFC"
GREEN_FILL = "DCFCE7"
AMBER_FILL = "FEF3C7"
RED_FILL = "FEE2E2"

PRIMARY = RGBColor(24, 38, 71)
SECONDARY = RGBColor(79, 98, 122)
ACCENT = RGBColor(193, 154, 107)
TEXT = RGBColor(0, 0, 0)
MUTED = RGBColor(102, 112, 122)
WHITE = RGBColor(255, 255, 255)
GREEN = RGBColor(22, 101, 52)
AMBER = RGBColor(180, 83, 9)
RED = RGBColor(185, 28, 28)
RAG_BANDS = [
    (0, 4, "#dc2626"),
    (4, 7, "#d97706"),
    (7, 10.01, "#16a34a"),
]

SYSTEM_PROMPT = """
You are writing a final client report for a Recruitment Operating Model Audit.

Write in British English. Use short sentences. Be direct, commercial and precise.

Rules:
- No markdown.
- No headings inside body text.
- No asterisks, hashes or raw bullet symbols.
- No filler or generic corporate language.
- Do not use these words or phrases: robust, hampered, material, leverage, seamless, best-in-class, best in class.
- Avoid repeating the company name unless needed.
- Keep the tone credible and board-ready.
- Keep the executive overview under 150 words.
- For each detailed section, keep current state to 3 sentences, key risks to 2 points, commercial impact to 2 sentences, immediate actions to 2 points, and next phase actions to 2 points.
""".strip()

REPORT_SCHEMA = {
    "type": "object",
    "properties": {
        "executive_overview": {"type": "string"},
        "top_strengths": {
            "type": "array",
            "items": {"type": "string"},
            "minItems": 5,
            "maxItems": 5,
        },
        "top_problems": {
            "type": "array",
            "items": {"type": "string"},
            "minItems": 5,
            "maxItems": 5,
        },
        "day_30_plan": {
            "type": "array",
            "items": {"type": "string"},
            "minItems": 3,
            "maxItems": 5,
        },
        "day_60_plan": {
            "type": "array",
            "items": {"type": "string"},
            "minItems": 3,
            "maxItems": 5,
        },
        "day_90_plan": {
            "type": "array",
            "items": {"type": "string"},
            "minItems": 3,
            "maxItems": 5,
        },
        "overall_recruitment_score": {"type": "string"},
        "final_verdict": {"type": "string"},
        "sections": {
            "type": "object",
            "properties": {
                section_id: {
                    "type": "object",
                    "properties": {
                        "score": {"type": "integer", "minimum": 1, "maximum": 10},
                        "current_state": {"type": "array", "items": {"type": "string"}, "minItems": 2, "maxItems": 4},
                        "key_risks": {"type": "array", "items": {"type": "string"}, "minItems": 2, "maxItems": 4},
                        "commercial_impact": {"type": "array", "items": {"type": "string"}, "minItems": 1, "maxItems": 3},
                        "immediate_actions": {"type": "array", "items": {"type": "string"}, "minItems": 2, "maxItems": 4},
                        "structural_improvements": {"type": "array", "items": {"type": "string"}, "minItems": 2, "maxItems": 4},
                    },
                    "required": [
                        "score",
                        "current_state",
                        "key_risks",
                        "commercial_impact",
                        "immediate_actions",
                        "structural_improvements",
                    ],
                    "additionalProperties": False,
                }
                for section_id in SECTION_IDS
            },
            "required": SECTION_IDS,
            "additionalProperties": False,
        },
    },
    "required": [
        "executive_overview",
        "top_strengths",
        "top_problems",
        "day_30_plan",
        "day_60_plan",
        "day_90_plan",
        "overall_recruitment_score",
        "final_verdict",
        "sections",
    ],
    "additionalProperties": False,
}


def get_api_key() -> str:
    candidates = [
        os.environ.get("OPENAI_API_KEY", ""),
        _read_text_if_exists(BASE_DIR / "openai_api_key.txt"),
        _read_text_if_exists(Path.home() / ".openai_api_key.txt"),
    ]
    for value in candidates:
        value = value.strip()
        if value:
            return value
    raise ValueError("OpenAI API key not found. Set OPENAI_API_KEY or add openai_api_key.txt alongside the app.")


def _read_text_if_exists(path: Path) -> str:
    return path.read_text(encoding="utf-8").strip() if path.exists() else ""


def parse_numeric_value(value: str | None) -> float | None:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    match = re.search(r"-?\d+(?:\.\d+)?", text.replace(",", ""))
    return float(match.group()) if match else None


def parse_time_to_hire_days(value: str | None) -> float | None:
    if value is None:
        return None
    text = str(value).strip().lower()
    if not text:
        return None
    numeric = parse_numeric_value(text)
    if numeric is None:
        return None
    if "week" in text or re.search(r"\bw\b", text):
        return round(numeric * 7, 1)
    if "month" in text:
        return round(numeric * 30, 1)
    return round(numeric, 1)


def list_benchmark_sectors() -> list[str]:
    df = _load_benchmark_table()
    return (
        df["sector"]
        .dropna()
        .astype(str)
        .str.strip()
        .replace("", pd.NA)
        .dropna()
        .drop_duplicates()
        .sort_values()
        .tolist()
    )


def load_benchmarks(sector: str) -> pd.DataFrame:
    df = _load_benchmark_table()

    if sector.strip():
        matches = df[df["sector"].str.contains(sector.strip(), case=False, na=False)]
        if not matches.empty:
            return matches.reset_index(drop=True)

    national = df[df["region"].astype(str).str.contains("UK National", case=False, na=False)]
    return national.reset_index(drop=True) if not national.empty else df.reset_index(drop=True)


def _load_benchmark_table() -> pd.DataFrame:
    benchmark_path = _resolve_benchmark_file()
    if benchmark_path.suffix.lower() == ".csv":
        df = pd.read_csv(benchmark_path)
    else:
        df = pd.read_excel(benchmark_path, sheet_name="Benchmarks")
    return _normalise_benchmark_columns(df)


def _resolve_benchmark_file() -> Path:
    candidates = []
    configured = os.environ.get(BENCHMARK_ENV_VAR, "").strip()
    if configured:
        candidates.append(Path(configured).expanduser())

    candidates.extend(
        [
            BENCHMARK_FILE,
            BENCHMARK_CSV_FILE,
            Path.cwd() / "uk_recruitment_benchmark_framework.xlsx",
            Path.cwd() / "uk_recruitment_benchmarks.csv",
            Path.home() / "Desktop" / "uk_recruitment_benchmarks.csv",
        ]
    )

    for path in candidates:
        if path.exists():
            return path

    searched = ", ".join(str(path) for path in candidates)
    raise FileNotFoundError(f"Benchmark dataset not found. Checked: {searched}")


def _normalise_benchmark_columns(df: pd.DataFrame) -> pd.DataFrame:
    renamed = df.copy()
    renamed.columns = [str(col).strip() for col in renamed.columns]
    renamed = renamed.rename(
        columns={
            "avg_time_to_hire": "avg_time_to_hire_days",
            "avg_applications": "avg_applications_per_role",
            "avg_offer_acceptance": "avg_offer_acceptance_pct",
            "avg_attrition": "avg_attrition_pct",
            "avg_application_to_interview": "avg_application_to_interview_pct",
            "avg_interview_to_offer": "avg_interview_to_offer_pct",
        }
    )

    if "sector" not in renamed.columns:
        raise ValueError("Benchmark data must include a 'sector' column.")

    defaults = {
        "company_size_band": "All sizes",
        "region": "UK National",
        "avg_time_to_hire_days": pd.NA,
        "avg_time_to_fill_days": pd.NA,
        "avg_applications_per_role": pd.NA,
        "avg_offer_acceptance_pct": pd.NA,
        "avg_attrition_pct": pd.NA,
        "avg_application_to_interview_pct": pd.NA,
        "avg_interview_to_offer_pct": pd.NA,
        "salary_competitiveness_index": pd.NA,
        "data_quality_note": "Imported benchmark dataset.",
        "source_basis": "Benchmark import",
    }
    for column, default in defaults.items():
        if column not in renamed.columns:
            renamed[column] = default

    renamed["sector"] = renamed["sector"].astype(str).str.strip()
    renamed["company_size_band"] = (
        renamed["company_size_band"].fillna("All sizes").astype(str).str.strip().replace("", "All sizes")
    )
    renamed["region"] = (
        renamed["region"].fillna("UK National").astype(str).str.strip().replace("", "UK National")
    )
    return renamed


def build_benchmark_summary(metrics: dict, benchmark: pd.DataFrame) -> dict:
    benchmark_row = _pick_benchmark_row(benchmark, {})
    if benchmark_row.empty:
        return {"benchmark_row": {}, "comparisons": [], "summary_text": "No benchmark data available."}

    mapping = [
        ("time_to_hire_days", "avg_time_to_hire_days", "Time to hire", "days", False),
        ("applications_per_role", "avg_applications_per_role", "Applications per role", "", True),
        ("offer_acceptance", "avg_offer_acceptance_pct", "Offer acceptance", "%", True),
        ("first_year_attrition", "avg_attrition_pct", "First-year attrition", "%", False),
    ]

    comparisons = []
    summary_lines = []
    for client_key, benchmark_key, label, suffix, higher_is_better in mapping:
        client_value = metrics.get(client_key)
        benchmark_value = _safe_float(benchmark_row.get(benchmark_key))
        if client_value is None or benchmark_value is None:
            continue

        delta = client_value - benchmark_value
        tolerance = max(1.0, abs(benchmark_value) * 0.03)
        if abs(delta) <= tolerance:
            status = "In line"
            comment = "In line with benchmark"
        else:
            ahead = delta > 0 if higher_is_better else delta < 0
            status = "Ahead" if ahead else "Behind"
            direction = "ahead of" if ahead else "behind"
            comment = f"{_format_metric_value(abs(delta), suffix)} {direction} benchmark"

        comparisons.append(
            {
                "label": label,
                "client_value": client_value,
                "benchmark_value": benchmark_value,
                "suffix": suffix,
                "status": status,
                "comment": comment,
            }
        )
        summary_lines.append(
            f"{label}: {_format_metric_value(client_value, suffix)} versus {_format_metric_value(benchmark_value, suffix)}. {comment}."
        )

    return {
        "benchmark_row": benchmark_row.to_dict(),
        "comparisons": comparisons,
        "summary_text": " ".join(summary_lines) if summary_lines else "No useful benchmark comparisons were available.",
    }


def auto_score_sections(data: dict, benchmark: pd.DataFrame) -> tuple[list[int], list[str]]:
    benchmark_row = _pick_benchmark_row(benchmark, data)
    metrics = data["metrics"]
    flags = data["process_flags"]

    metric_scores = {
        "time_to_hire": _metric_score(metrics.get("time_to_hire_days"), _safe_float(benchmark_row.get("avg_time_to_hire_days")), higher_is_better=False),
        "applications": _metric_score(metrics.get("applications_per_role"), _safe_float(benchmark_row.get("avg_applications_per_role")), higher_is_better=True),
        "offer_acceptance": _metric_score(metrics.get("offer_acceptance"), _safe_float(benchmark_row.get("avg_offer_acceptance_pct")), higher_is_better=True),
        "attrition": _metric_score(metrics.get("first_year_attrition"), _safe_float(benchmark_row.get("avg_attrition_pct")), higher_is_better=False),
        "feedback_speed": _feedback_score(metrics.get("interview_feedback_time_days")),
        "interview_design": _stage_score(metrics.get("interview_stages")),
        "screening_volume": _screening_score(metrics.get("candidates_reaching_interview")),
    }

    def flag_value(*names: str) -> float:
        if not names:
            return 5.0
        yes = sum(1 for name in names if flags.get(name))
        return 3.5 + (yes / len(names)) * 5.5

    def blend(*values: float) -> int:
        return max(1, min(10, round(sum(values) / len(values))))

    scores = [
        blend(flag_value("has_hiring_plan", "named_process_owner"), metric_scores["time_to_hire"], metric_scores["attrition"]),
        blend(flag_value("tracks_metrics"), metric_scores["offer_acceptance"], metric_scores["applications"], metric_scores["screening_volume"]),
        blend(flag_value("has_employer_brand"), metric_scores["applications"], metric_scores["offer_acceptance"]),
        blend(flag_value("standardised_job_specs"), metric_scores["applications"], metric_scores["offer_acceptance"]),
        blend(flag_value("multi_channel_sourcing"), metric_scores["applications"], metric_scores["screening_volume"]),
        blend(flag_value("structured_screening"), metric_scores["screening_volume"], metric_scores["feedback_speed"]),
        blend(flag_value("structured_interviews", "hiring_manager_training"), metric_scores["interview_design"], metric_scores["feedback_speed"]),
        blend(flag_value("fast_offer_process"), metric_scores["offer_acceptance"], metric_scores["feedback_speed"]),
        blend(flag_value("formal_onboarding"), metric_scores["attrition"], metric_scores["offer_acceptance"]),
        blend(metric_scores["attrition"], metric_scores["time_to_hire"]),
        blend(flag_value("collects_candidate_feedback"), metric_scores["feedback_speed"], metric_scores["offer_acceptance"]),
        blend(flag_value("named_process_owner", "tracks_metrics", "hiring_manager_training")),
    ]

    notes = [
        f"Planning maturity is {'stronger' if flags.get('has_hiring_plan') else 'weaker'} than ideal, with time to hire at {_fmt(metrics.get('time_to_hire_days'), ' days')}.",
        f"KPI visibility is {'embedded' if flags.get('tracks_metrics') else 'limited'}, with offer acceptance at {_fmt(metrics.get('offer_acceptance'), '%')}.",
        f"Employer brand is {'defined' if flags.get('has_employer_brand') else 'underdeveloped'}, and applications per role are {_fmt(metrics.get('applications_per_role'), '')}.",
        f"Job documentation is {'standardised' if flags.get('standardised_job_specs') else 'inconsistent'}, affecting role clarity and shortlist quality.",
        f"Sourcing mix is {'broad' if flags.get('multi_channel_sourcing') else 'narrow'}, with interview-ready candidate flow at {_fmt(metrics.get('candidates_reaching_interview'), '')}.",
        f"Screening is {'structured' if flags.get('structured_screening') else 'informal'}, which influences response speed and conversion quality.",
        f"Interview design combines {int(metrics.get('interview_stages') or 0) or 'an unknown number of'} stages with feedback turnaround of {_fmt(metrics.get('interview_feedback_time_days'), ' days')}.",
        f"Offer process is {'faster' if flags.get('fast_offer_process') else 'more exposed to delay'}, which shapes acceptance performance.",
        f"Onboarding is {'documented' if flags.get('formal_onboarding') else 'lightly controlled'}, with first-year attrition at {_fmt(metrics.get('first_year_attrition'), '%')}.",
        f"Retention pressure remains visible through first-year attrition of {_fmt(metrics.get('first_year_attrition'), '%')}.",
        f"Candidate feedback is {'captured' if flags.get('collects_candidate_feedback') else 'not systematically captured'}, while interview feedback speed still needs attention.",
        f"Ownership is {'clear' if flags.get('named_process_owner') else 'blurred'}, and manager training is {'present' if flags.get('hiring_manager_training') else 'limited'}.",
    ]

    return scores, notes


def generate_report_json(client, data: dict, benchmark_summary: dict) -> dict:
    prompt = _build_user_prompt(data, benchmark_summary)
    response = client.responses.create(
        model="gpt-4.1",
        instructions=SYSTEM_PROMPT,
        input=prompt,
        max_output_tokens=5000,
        temperature=0.5,
        text={
            "verbosity": "medium",
            "format": {
                "type": "json_schema",
                "name": "recruitment_audit_report",
                "schema": REPORT_SCHEMA,
                "strict": True,
            },
        },
    )
    report = json.loads(response.output_text)
    return _normalise_report(report, data["section_scores"])


def _build_user_prompt(data: dict, benchmark_summary: dict) -> str:
    strongest_sections = sorted(
        zip(SECTION_ORDER, data["section_scores"]),
        key=lambda item: item[1],
        reverse=True,
    )[:3]
    weakest_sections = sorted(
        zip(SECTION_ORDER, data["section_scores"]),
        key=lambda item: item[1],
    )[:3]
    section_scorecard = "\n".join(
        f"- {title}: {score}/10. {note}"
        for title, score, note in zip(SECTION_ORDER, data["section_scores"], data["section_notes"])
    )
    metrics_block = "\n".join(
        f"- {key.replace('_', ' ').title()}: {value or 'Not provided'}"
        for key, value in data["raw_metrics"].items()
    )
    flags_block = "\n".join(
        f"- {key.replace('_', ' ').title()}: {'Yes' if value else 'No'}"
        for key, value in data["process_flags"].items()
    )
    return f"""
Company profile
- Company: {data['company_name']}
- Contact name: {data['contact_name']}
- Job title: {data['job_title']}
- Phone number: {data['phone_number']}
- Email address: {data['email_address']}
- Office address: {data['office_address']}
- Sector: {data['sector']}
- Location: {data['location']}
- Headcount: {data['headcount']}
- Annual hiring volume: {data['annual_hiring_volume']}
- Key roles / job titles hired: {data['key_roles_hired']}

Metrics supplied
{metrics_block}

Operating controls
{flags_block}

Benchmark summary
{benchmark_summary['summary_text']}

Calculated scorecard
{section_scorecard}

Strongest areas
{chr(10).join(f"- {title}: {score}/10" for title, score in strongest_sections)}

Weakest areas
{chr(10).join(f"- {title}: {score}/10" for title, score in weakest_sections)}

Output requirements
- Use the same score for each detailed section as the supplied scorecard.
- Make the executive overview flow as a narrative, not as bullets.
- Keep strengths, problems and plans crisp and commercially grounded.
- Do not invent client facts that are not implied by the supplied data.
- Make the detailed sections feel connected rather than written in isolation.
- Where a benchmark gap is clear, say what it means operationally.
""".strip()


def create_section_score_chart(company_name: str, section_scores: list[int]) -> Path:
    output_dir = _output_dir()
    path = output_dir / f"{_slug(company_name)}_section_scores.png"

    labels = [_short_label(title) for title in SECTION_ORDER]
    positions = list(range(len(labels)))

    fig, ax = plt.subplots(figsize=(7.1, 4.8))
    _apply_chart_style(fig, ax)
    bar_colors = [_score_hex(score) for score in section_scores]
    ax.barh(positions, section_scores, color=bar_colors, edgecolor="black", linewidth=0.5, height=0.56)
    ax.set_xlim(0, 10)
    ax.set_xticks(range(0, 11, 2))
    ax.set_yticks(positions, labels)
    ax.invert_yaxis()
    ax.grid(axis="x", alpha=0.12, linestyle="--")
    ax.set_title("Section scores", fontsize=12.5, pad=10, color="black", fontweight="bold")
    ax.set_xlabel("Score out of 10", color="black")
    for pos, score in enumerate(section_scores):
        ax.text(min(score + 0.15, 9.7), pos, f"{score}/10", va="center", fontsize=8.4, color="black", fontweight="bold")
    fig.text(0.11, 0.95, company_name, fontsize=9, color="black")
    fig.tight_layout(rect=(0, 0, 1, 0.93))
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


def create_overall_score_chart(company_name: str, total_score: int) -> Path:
    output_dir = _output_dir()
    path = output_dir / f"{_slug(company_name)}_overall_score.png"

    max_score = 120
    score = min(total_score, max_score)
    pct = round((score / max_score) * 100)
    rating = _rating_for_score(score)
    fill_color = "#333333"

    fig, ax = plt.subplots(figsize=(7.0, 1.55))
    _apply_chart_style(fig, ax)
    ax.barh([0], [max_score], color="#E5E5E5", height=0.28)
    ax.barh([0], [score], color=fill_color, height=0.34)
    ax.set_xlim(0, max_score)
    ax.set_ylim(-0.42, 0.42)
    ax.set_yticks([])
    ax.set_xticks(range(0, max_score + 1, 20))
    ax.tick_params(axis="x", labelsize=8.2, colors="black")
    ax.grid(axis="x", alpha=0.12, linestyle="--")
    fig.text(0.1, 0.92, company_name, fontsize=9, color="black")
    fig.text(0.1, 0.82, "Overall score", fontsize=12.5, color="black", fontweight="bold")
    fig.text(0.9, 0.84, f"{score}/120 | {pct}% | {rating}", ha="right", fontsize=9.5, color="black", fontweight="bold")
    ax.annotate(
        f"{score}/120 ({pct}%)",
        xy=(score, 0),
        xytext=(0, 10),
        textcoords="offset points",
        ha="center",
        va="bottom",
        fontsize=8.5,
        color="black",
        fontweight="bold",
    )
    for spine in ax.spines.values():
        spine.set_visible(False)
    fig.subplots_adjust(left=0.09, right=0.97, bottom=0.24, top=0.62)
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


def create_benchmark_chart(company_name: str, metrics: dict, benchmark: pd.DataFrame) -> Path:
    output_dir = _output_dir()
    path = output_dir / f"{_slug(company_name)}_benchmark_compare.png"
    benchmark_row = _pick_benchmark_row(benchmark, {})

    items = [
        ("Time to hire", metrics.get("time_to_hire_days"), _safe_float(benchmark_row.get("avg_time_to_hire_days")), "days", False),
        ("Applications per role", metrics.get("applications_per_role"), _safe_float(benchmark_row.get("avg_applications_per_role")), "", True),
        ("Offer acceptance", metrics.get("offer_acceptance"), _safe_float(benchmark_row.get("avg_offer_acceptance_pct")), "%", True),
        ("First-year attrition", metrics.get("first_year_attrition"), _safe_float(benchmark_row.get("avg_attrition_pct")), "%", False),
    ]
    items = [item for item in items if item[1] is not None and item[2] is not None]

    if not items:
        fig, ax = plt.subplots(figsize=(7.0, 1.5))
        _apply_chart_style(fig, ax)
        ax.axis("off")
        ax.text(0.5, 0.5, "No benchmark comparison available", ha="center", va="center", fontsize=10, color="black")
        fig.savefig(path, dpi=220, bbox_inches="tight", facecolor=fig.get_facecolor())
        plt.close(fig)
        return path

    fig, axes = plt.subplots(len(items), 1, figsize=(7.0, 1.2 + (len(items) * 1.0)))
    if len(items) == 1:
        axes = [axes]
    for ax, (label, client_value, benchmark_value, suffix, higher_is_better) in zip(axes, items):
        _apply_chart_style(fig, ax)
        lower = min(client_value, benchmark_value, 0)
        upper = max(client_value, benchmark_value, 1)
        spread = max(upper - lower, upper * 0.2, 1)
        x_min = max(0, lower - spread * 0.2)
        x_max = upper + spread * 0.2
        ahead = client_value >= benchmark_value if higher_is_better else client_value <= benchmark_value
        client_color = "#222222"
        ax.set_xlim(x_min, x_max)
        ax.set_ylim(-0.5, 0.5)
        ax.hlines(0, x_min, x_max, color="#BFBFBF", linewidth=1.4)
        ax.scatter([benchmark_value], [0], s=60, color="#D7D7D7", edgecolors="black", linewidths=0.6, zorder=3)
        ax.scatter([client_value], [0], s=65, color=client_color, edgecolors="black", linewidths=0.6, marker="D", zorder=4)
        ax.set_yticks([])
        ax.xaxis.set_major_locator(MaxNLocator(4))
        ax.tick_params(axis="x", labelsize=8, colors="black")
        ax.set_title(label, fontsize=10.2, color="black", fontweight="bold", pad=2)
        ax.grid(axis="x", alpha=0.12, linestyle="--")
        client_offset = 10 if client_value >= benchmark_value else -12
        client_va = "bottom" if client_value >= benchmark_value else "top"
        benchmark_offset = -12 if client_value >= benchmark_value else 10
        benchmark_va = "top" if client_value >= benchmark_value else "bottom"
        ax.annotate(
            f"Client {_format_metric_value(client_value, suffix)}",
            xy=(client_value, 0),
            xytext=(0, client_offset),
            textcoords="offset points",
            ha="center",
            va=client_va,
            fontsize=8.1,
            color="black",
            fontweight="bold",
        )
        ax.annotate(
            f"Benchmark {_format_metric_value(benchmark_value, suffix)}",
            xy=(benchmark_value, 0),
            xytext=(0, benchmark_offset),
            textcoords="offset points",
            ha="center",
            va=benchmark_va,
            fontsize=8,
            color="black",
        )
        delta = abs(client_value - benchmark_value)
        direction = "Ahead of benchmark" if ahead else "Behind benchmark"
        ax.text(
            0.0,
            0.04,
            f"{direction} by {_format_metric_value(delta, suffix)}",
            transform=ax.transAxes,
            fontsize=8,
            color="black",
            fontweight="bold",
        )
        for spine in ax.spines.values():
            spine.set_visible(False)

    fig.suptitle("Benchmark comparison", fontsize=12.5, y=0.99, color="black", fontweight="bold")
    fig.text(0.11, 0.95, company_name, fontsize=9, color="black")
    fig.tight_layout(rect=(0, 0, 1, 0.93), h_pad=0.7)
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor=fig.get_facecolor())
    plt.close(fig)
    return path


def save_word_report(
    data: dict,
    report: dict,
    benchmark_summary: dict,
    section_chart: Path,
    overall_chart: Path,
    benchmark_chart: Path,
) -> Path:
    output_path = _output_dir() / f"{_slug(data['company_name'])}_recruitment_audit.docx"
    report = _clean_report(_normalise_report(report, data["section_scores"]), data)
    document = Document()
    _set_document_defaults(document)
    _add_cover_page(document, data)
    _add_executive_overview(document, report)
    _add_overall_score_section(document, data, report, overall_chart)
    _add_key_insights(document, data, report, benchmark_summary)
    _add_score_summary(document, data)
    _add_benchmark_snapshot(document, benchmark_summary)
    _add_priority_matrix(document, data, report)
    _add_chart_section(document, [section_chart, benchmark_chart])
    document.add_page_break()
    _add_detailed_findings(document, data, report)
    document.add_page_break()
    _add_roadmap(document, report)
    _add_final_verdict(document, report)
    document.save(output_path)
    return output_path


def _normalise_report(report: dict, fallback_scores: list[int]) -> dict:
    sections_value = report.get("sections", {}) or {}
    if isinstance(sections_value, list):
        normalised_sections = []
        for index, section in enumerate(sections_value):
            section = section or {}
            normalised_sections.append(
                {
                    "title": str(section.get("title", SECTION_ORDER[index])),
                    "score": int(section.get("score", fallback_scores[index])),
                    "current_state": _ensure_list(section.get("current_state"), 2),
                    "key_risks": _ensure_list(section.get("key_risks"), 2),
                    "commercial_impact": _ensure_list(section.get("commercial_impact"), 1),
                    "immediate_actions": _ensure_list(section.get("immediate_actions"), 2),
                    "structural_improvements": _ensure_list(section.get("structural_improvements"), 2),
                }
            )
        report["sections"] = normalised_sections
        return report

    sections_by_id = sections_value
    sections = []
    for index, section_id in enumerate(SECTION_IDS):
        title = SECTION_ID_TO_TITLE[section_id]
        section = sections_by_id.get(section_id, {})
        sections.append(
            {
                "title": title,
                "score": int(section.get("score", fallback_scores[index])),
                "current_state": _ensure_list(section.get("current_state"), 2),
                "key_risks": _ensure_list(section.get("key_risks"), 2),
                "commercial_impact": _ensure_list(section.get("commercial_impact"), 1),
                "immediate_actions": _ensure_list(section.get("immediate_actions"), 2),
                "structural_improvements": _ensure_list(section.get("structural_improvements"), 2),
            }
        )
    report["sections"] = sections
    return report


def _clean_report(report: dict, data: dict) -> dict:
    cleaned = {
        "executive_overview": _clean_text(report.get("executive_overview"), max_sentences=6, max_words=150),
        "top_strengths": _clean_list(report.get("top_strengths"), max_items=5, max_words=18),
        "top_problems": _clean_list(report.get("top_problems"), max_items=5, max_words=18),
        "day_30_plan": _clean_list(report.get("day_30_plan"), max_items=5, max_words=18),
        "day_60_plan": _clean_list(report.get("day_60_plan"), max_items=5, max_words=18),
        "day_90_plan": _clean_list(report.get("day_90_plan"), max_items=5, max_words=18),
        "overall_recruitment_score": _clean_text(report.get("overall_recruitment_score"), max_sentences=2, max_words=45),
        "final_verdict": _clean_text(report.get("final_verdict"), max_sentences=3, max_words=70),
        "sections": [],
    }

    for index, section in enumerate(report.get("sections", [])):
        section = section or {}
        cleaned["sections"].append(
            {
                "title": SECTION_ORDER[index],
                "score": int(section.get("score", data["section_scores"][index])),
                "current_state": _compose_current_state(section.get("current_state"), data["section_notes"][index]),
                "key_risks": _clean_list(section.get("key_risks"), max_items=2, max_words=18),
                "commercial_impact": _clean_text(section.get("commercial_impact"), max_sentences=2, max_words=38),
                "immediate_actions": _clean_list(section.get("immediate_actions"), max_items=2, max_words=18),
                "structural_improvements": _clean_list(section.get("structural_improvements"), max_items=2, max_words=18),
            }
        )

    if not cleaned["executive_overview"]:
        cleaned["executive_overview"] = (
            "This audit shows how the recruitment function is performing in practice. "
            "The current operating model has clear strengths, but control and pace are inconsistent in weaker areas. "
            "The priority is to improve the parts of the process that are slowing hiring or reducing decision quality."
        )
    if not cleaned["overall_recruitment_score"]:
        cleaned["overall_recruitment_score"] = (
            "The overall score points to a recruitment function that is workable, but not yet consistent enough in the areas that matter most."
        )
    if not cleaned["final_verdict"]:
        cleaned["final_verdict"] = (
            "The recruitment function can support current hiring demand, but it needs tighter control in weaker areas to improve pace, consistency and confidence."
        )
    return cleaned


def _ensure_list(value, min_items: int) -> list[str]:
    if isinstance(value, list):
        items = [str(item).strip() for item in value if str(item).strip()]
    elif value:
        items = [str(value).strip()]
    else:
        items = []
    while len(items) < min_items:
        items.append("This area needs clearer definition in the next review cycle.")
    return items


def _set_document_defaults(document: Document) -> None:
    font_name = "Aptos"
    normal = document.styles["Normal"]
    normal.font.name = font_name
    normal.font.size = Pt(10)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)

    heading_specs = {
        "Title": 22,
        "Heading 1": 14,
        "Heading 2": 11.5,
        "Heading 3": 10,
    }
    for name, size in heading_specs.items():
        if name not in document.styles:
            continue
        style = document.styles[name]
        style.font.name = font_name
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = TEXT

    if "List Bullet" in document.styles:
        style = document.styles["List Bullet"]
        style.font.name = font_name
        style.font.size = Pt(10)
        style.font.color.rgb = TEXT

    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def _add_cover_page(document: Document, data: dict) -> None:
    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(125)
    title.paragraph_format.space_after = Pt(8)
    title.style = document.styles["Title"]
    title.add_run(REPORT_TITLE)

    company = document.add_paragraph()
    company.paragraph_format.space_after = Pt(18)
    run = company.add_run(data["company_name"].strip())
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(15)
    run.font.color.rgb = TEXT

    detail_lines = [
        f"Report date: {datetime.now().strftime('%d %B %Y')}",
        f"Sector: {data['sector']}",
        f"Location: {data['location']}",
        f"Headcount: {data['headcount']}",
        f"Annual hiring volume: {data['annual_hiring_volume']}",
        f"Contact: {data['contact_name']} | {data['job_title']}",
        f"Email: {data['email_address']}",
        f"Phone: {data['phone_number']}",
        f"Office address: {data['office_address']}",
    ]
    for line in detail_lines:
        _add_paragraph(document, line, after=2)
    document.add_page_break()


def _add_executive_overview(document: Document, report: dict) -> None:
    _add_heading(document, "Executive overview", level=1)
    _add_paragraph(document, report["executive_overview"], after=8)


def _add_overall_score_section(document: Document, data: dict, report: dict, chart_path: Path) -> None:
    _add_heading(document, "Overall score", level=1)

    score = int(data["total_score"])
    percentage = round((score / 120) * 100)
    rating = _rating_for_score(score)

    table = document.add_table(rows=2, cols=4)
    table.autofit = True
    labels = ["Total score", "Percentage", "Rating", "Key roles"]
    values = [f"{score}/120", f"{percentage}%", rating, data["key_roles_hired"]]
    for index, label in enumerate(labels):
        _shade_cell(table.cell(0, index), "EDEDED")
        _set_cell(table.cell(0, index), label, bold=True)
        _set_cell(table.cell(1, index), values[index])

    document.add_paragraph("")
    if chart_path.exists():
        document.add_picture(str(chart_path), width=Inches(5.4))
    _add_paragraph(document, report["overall_recruitment_score"], after=8)


def _add_key_insights(document: Document, data: dict, report: dict, benchmark_summary: dict) -> None:
    _add_heading(document, "Key insights", level=1)
    for line in _build_key_insights(data, report, benchmark_summary):
        _add_paragraph(document, line, after=2)
    document.add_paragraph("")


def _add_score_summary(document: Document, data: dict) -> None:
    _add_heading(document, "Score summary", level=1)
    table = document.add_table(rows=1, cols=3)
    table.autofit = True
    header = table.rows[0].cells
    for index, label in enumerate(["Area", "Score", "Rating"]):
        _shade_cell(header[index], "EDEDED")
        _set_cell(header[index], label, bold=True)

    for title, score in zip(SECTION_ORDER, data["section_scores"]):
        row = table.add_row().cells
        _set_cell(row[0], title)
        _set_cell(row[1], f"{score}/10", bold=True)
        _set_cell(row[2], _section_rating(score))
    document.add_paragraph("")


def _add_benchmark_snapshot(document: Document, benchmark_summary: dict) -> None:
    comparisons = benchmark_summary.get("comparisons", [])
    if not comparisons:
        return

    _add_heading(document, "Benchmark snapshot", level=1)
    table = document.add_table(rows=1, cols=4)
    table.autofit = True
    header = table.rows[0].cells
    for idx, label in enumerate(["Metric", "Client", "Benchmark", "Comment"]):
        _shade_cell(header[idx], "EDEDED")
        _set_cell(header[idx], label, bold=True)

    for item in comparisons:
        cells = table.add_row().cells
        suffix = item["suffix"]
        _set_cell(cells[0], item["label"])
        _set_cell(cells[1], _format_metric_value(item["client_value"], suffix))
        _set_cell(cells[2], _format_metric_value(item["benchmark_value"], suffix))
        _shade_cell(cells[3], _status_fill(item["status"]))
        _set_cell(cells[3], item["comment"])
    document.add_paragraph("")


def _add_priority_matrix(document: Document, data: dict, report: dict) -> None:
    _add_heading(document, "Priority matrix", level=1)
    _add_paragraph(
        document,
        "These priorities focus on the areas with the weakest control, the clearest commercial drag and the strongest case for action.",
        after=4,
    )

    priorities = _build_priority_matrix(data, report)
    table = document.add_table(rows=1, cols=4)
    table.autofit = True
    header = table.rows[0].cells
    for idx, label in enumerate(["Priority area", "Why it matters", "Action", "Timing"]):
        _shade_cell(header[idx], "EDEDED")
        _set_cell(header[idx], label, bold=True)

    for item in priorities:
        cells = table.add_row().cells
        _set_cell(cells[0], item["title"])
        _set_cell(cells[1], item["why"])
        _set_cell(cells[2], item["action"])
        _set_cell(cells[3], item["urgency"])
    document.add_paragraph("")


def _add_chart_section(document: Document, chart_paths: list[Path]) -> None:
    _add_heading(document, "Charts", level=1)
    _add_paragraph(
        document,
        "The charts below support the score and benchmark findings. They are scaled for board-level review.",
        after=6,
    )
    captions = [
        "Section-by-section scoring profile.",
        "Benchmark comparison for the most useful submitted metrics.",
    ]
    for chart_path, caption in zip(chart_paths, captions):
        if chart_path.exists():
            document.add_picture(str(chart_path), width=Inches(5.4))
            _add_paragraph(document, caption, after=8)
            document.add_paragraph("")


def _add_detailed_findings(document: Document, data: dict, report: dict) -> None:
    _add_heading(document, "Detailed findings", level=1)
    for section in report["sections"]:
        _add_heading(document, section["title"], level=2)
        _add_metric_callout(document, f"Score: {section['score']}/10")
        _add_heading(document, "Current state", level=3)
        _add_paragraph(document, section["current_state"], after=3)
        _add_heading(document, "Key risks", level=3)
        _add_bullets(document, section["key_risks"])
        _add_heading(document, "Commercial impact", level=3)
        _add_paragraph(document, section["commercial_impact"], after=3)
        _add_heading(document, "Actions", level=3)
        _add_action_line(document, "Immediate", section["immediate_actions"])
        _add_action_line(document, "Next phase", section["structural_improvements"])
        document.add_paragraph("")


def _add_roadmap(document: Document, report: dict) -> None:
    _add_heading(document, "Top 5 strengths", level=1)
    _add_bullets(document, report["top_strengths"])
    document.add_paragraph("")

    _add_heading(document, "Top 5 problems", level=1)
    _add_bullets(document, report["top_problems"])
    document.add_paragraph("")

    _add_heading(document, "30 day plan", level=1)
    _add_bullets(document, report["day_30_plan"])
    document.add_paragraph("")

    _add_heading(document, "60 day plan", level=1)
    _add_bullets(document, report["day_60_plan"])
    document.add_paragraph("")

    _add_heading(document, "90 day plan", level=1)
    _add_bullets(document, report["day_90_plan"])
    document.add_paragraph("")


def _add_final_verdict(document: Document, report: dict) -> None:
    _add_heading(document, "Final verdict", level=1)
    _add_paragraph(document, report["final_verdict"], after=0)


def _add_heading(document: Document, title: str, level: int) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = document.styles[f"Heading {min(level, 3)}"]
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8 if level == 2 else 4)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(14 if level == 1 else 11.5 if level == 2 else 10)
    run.font.color.rgb = TEXT


def _add_paragraph(document: Document, text: str, after: float = 3) -> None:
    text = _clean_text(text)
    if not text:
        return
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text.strip())
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT


def _add_bullets(document: Document, items: list[str]) -> None:
    for item in _clean_list(items, max_items=len(items) if items else 0, max_words=18):
        bullet = document.add_paragraph(style="List Bullet")
        bullet.paragraph_format.space_before = Pt(0)
        bullet.paragraph_format.space_after = Pt(1)
        bullet.paragraph_format.left_indent = Inches(0.18)
        run = bullet.add_run(item)
        run.font.name = "Aptos"
        run.font.size = Pt(10)
        run.font.color.rgb = TEXT


def _add_action_line(document: Document, label: str, items: list[str]) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.1
    lead = paragraph.add_run(f"{label}: ")
    lead.bold = True
    lead.font.name = "Aptos"
    lead.font.size = Pt(10)
    lead.font.color.rgb = TEXT
    body = paragraph.add_run("; ".join(_clean_list(items, max_items=2, max_words=18)))
    body.font.name = "Aptos"
    body.font.size = Pt(10)
    body.font.color.rgb = TEXT


def _add_metric_callout(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(10)
    run.font.color.rgb = TEXT


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell(
    cell,
    text: str,
    bold: bool = False,
    color: RGBColor = TEXT,
    size: float = 10.0,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = alignment
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run(_clean_text(text))
    run.bold = bold
    run.font.name = "Aptos"
    run.font.size = Pt(size)
    run.font.color.rgb = TEXT
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _apply_chart_style(fig, ax) -> None:
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    ax.tick_params(axis="x", colors="black", labelsize=8.5)
    ax.tick_params(axis="y", colors="black", labelsize=8.5)
    for spine in ax.spines.values():
        spine.set_visible(False)


def _output_dir() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    return OUTPUT_DIR


def _slug(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    cleaned = re.sub(r"_+", "_", cleaned).strip("._")
    return cleaned or "recruitment_audit"


def _size_band(headcount_value: str) -> str:
    headcount = parse_numeric_value(headcount_value)
    if headcount is None:
        return "SME"
    if headcount >= 1000:
        return "Enterprise"
    if headcount >= 250:
        return "Mid-Market"
    return "SME"


def _region_hint(location: str) -> str:
    text = location.lower()
    mapping = {
        "london": "London",
        "midlands": "Midlands",
        "manchester": "North West",
        "liverpool": "North West",
        "north west": "North West",
        "scotland": "Scotland",
        "glasgow": "Scotland",
        "edinburgh": "Scotland",
        "surrey": "South East",
        "kent": "South East",
        "south east": "South East",
    }
    for needle, region in mapping.items():
        if needle in text:
            return region
    return "UK National"


def _pick_benchmark_row(benchmark: pd.DataFrame, data: dict) -> pd.Series:
    if benchmark.empty:
        return pd.Series(dtype=object)

    working = benchmark.copy()
    size_band = _size_band(str(data.get("headcount", "")))
    region = _region_hint(str(data.get("location", "")))

    sized = working[working["company_size_band"].astype(str) == size_band]
    if not sized.empty:
        working = sized

    regional = working[working["region"].astype(str) == region]
    if regional.empty:
        regional = working[working["region"].astype(str) == "UK National"]
    if not regional.empty:
        working = regional

    return working.iloc[0]


def _metric_score(value: float | None, benchmark_value: float | None, higher_is_better: bool) -> float:
    if value is None or benchmark_value is None or benchmark_value == 0:
        return 5.5
    ratio = value / benchmark_value
    if higher_is_better:
        if ratio >= 1.2:
            return 8.8
        if ratio >= 1.05:
            return 7.8
        if ratio >= 0.95:
            return 6.8
        if ratio >= 0.8:
            return 5.5
        return 4.0
    if ratio <= 0.8:
        return 8.8
    if ratio <= 0.95:
        return 7.8
    if ratio <= 1.05:
        return 6.8
    if ratio <= 1.2:
        return 5.5
    return 4.0


def _feedback_score(days: float | None) -> float:
    if days is None:
        return 5.5
    if days <= 2:
        return 9.0
    if days <= 4:
        return 7.5
    if days <= 7:
        return 5.5
    return 3.8


def _stage_score(stages: float | None) -> float:
    if stages is None:
        return 5.5
    if stages <= 2:
        return 8.8
    if stages <= 3:
        return 7.2
    if stages <= 4:
        return 5.2
    return 3.8


def _screening_score(candidates: float | None) -> float:
    if candidates is None:
        return 5.5
    if candidates >= 5:
        return 8.0
    if candidates >= 3:
        return 6.8
    if candidates >= 2:
        return 5.4
    return 4.0


def _safe_float(value) -> float | None:
    if value is None or pd.isna(value):
        return None
    return float(value)


def _short_label(title: str) -> str:
    mapping = {
        "Recruitment strategy and workforce planning": "Strategy & planning",
        "Performance metrics and funnel conversion": "Metrics & funnel",
        "Employer brand and market perception": "Employer brand",
        "Job adverts and job specifications": "Job adverts & specs",
        "Sourcing and advertising process": "Sourcing process",
        "Application handling and screening": "Screening",
        "Interview process quality": "Interviews",
        "Decision making and offer process": "Offers",
        "Onboarding and early retention": "Onboarding",
        "Staff turnover risks": "Turnover risk",
        "Candidate experience": "Candidate experience",
        "Process ownership and accountability": "Ownership",
    }
    return "\n".join(textwrap.wrap(mapping.get(title, title), width=18))


def _rating_for_score(total_score: int) -> str:
    if total_score >= 96:
        return "High maturity"
    if total_score >= 78:
        return "Established with gaps"
    if total_score >= 60:
        return "Functional but inconsistent"
    return "Needs substantial improvement"


def _score_rgb(score: int) -> RGBColor:
    return TEXT


def _score_fill(score: int) -> str:
    if score >= 7:
        return "D9D9D9"
    if score >= 4:
        return "EBEBEB"
    return "F5F5F5"


def _score_hex(score: int) -> str:
    if score >= 7:
        return "#3A3A3A"
    if score >= 4:
        return "#6A6A6A"
    return "#9A9A9A"


def _status_colors(status: str) -> tuple[RGBColor, str]:
    return TEXT, _status_fill(status)


def _fmt(value: float | None, suffix: str) -> str:
    return "not provided" if value is None else _format_metric_value(value, suffix)


def _build_priority_matrix(data: dict, report: dict) -> list[dict[str, str]]:
    priorities = []
    sections = sorted(report.get("sections", []), key=lambda item: item["score"])[:4]
    for section in sections:
        urgency = "Immediate" if section["score"] <= 4 else "Next 30 days" if section["score"] <= 6 else "Planned"
        priorities.append(
            {
                "title": section["title"],
                "urgency": urgency,
                "impact": _section_rating(section["score"]),
                "why": section["commercial_impact"],
                "action": section["immediate_actions"][0],
            }
        )
    return priorities


def _clean_text(value, max_sentences: int | None = None, max_words: int | None = None) -> str:
    if value is None:
        return ""
    if isinstance(value, list):
        value = " ".join(str(item) for item in value)
    text = str(value)
    text = text.replace("\r", "\n")
    text = re.sub(r"(?m)^\s*[*#>-]+\s*", "", text)
    text = re.sub(r"(?m)^\s*\d+[.)]\s*", "", text)
    text = text.replace("•", " ")
    text = text.replace("–", "-").replace("—", "-")
    text = re.sub(r"\s+", " ", text).strip()

    replacements = {
        "best-in-class": "strong",
        "best in class": "strong",
        "robust": "clear",
        "hampered": "slowed",
        "material": "clear",
        "leverage": "use",
        "seamless": "consistent",
    }
    for source, target in replacements.items():
        text = re.sub(rf"\b{re.escape(source)}\b", target, text, flags=re.IGNORECASE)

    text = re.sub(r"\s+([,.;:])", r"\1", text)
    text = re.sub(r"([.?!]){2,}", r"\1", text)

    if max_sentences is not None:
        sentences = _split_sentences(text)
        text = " ".join(sentences[:max_sentences])
    if max_words is not None:
        words = text.split()
        if len(words) > max_words:
            text = " ".join(words[:max_words]).rstrip(",;:")
            if text and text[-1] not in ".!?":
                text += "."
    return text.strip()


def _clean_list(items, max_items: int, max_words: int = 18) -> list[str]:
    if items is None:
        source_items = []
    elif isinstance(items, list):
        source_items = items
    else:
        source_items = [items]

    cleaned = []
    for item in source_items:
        text = _clean_text(item, max_sentences=1, max_words=max_words)
        if text:
            cleaned.append(text)
        if max_items and len(cleaned) >= max_items:
            break
    return cleaned


def _compose_current_state(value, fallback_note: str) -> str:
    sentences = []
    for text in _clean_list(value, max_items=3, max_words=22):
        sentences.extend(_split_sentences(text))
    if fallback_note:
        sentences.extend(_split_sentences(_clean_text(fallback_note, max_sentences=2, max_words=28)))

    unique = []
    seen = set()
    for sentence in sentences:
        key = sentence.lower()
        if key not in seen:
            unique.append(sentence)
            seen.add(key)
        if len(unique) == 3:
            break
    return " ".join(unique)


def _split_sentences(text: str) -> list[str]:
    text = text.strip()
    if not text:
        return []
    parts = re.split(r"(?<=[.!?])\s+", text)
    cleaned = [part.strip() for part in parts if part.strip()]
    return cleaned or [text]


def _build_key_insights(data: dict, report: dict, benchmark_summary: dict) -> list[str]:
    strongest = max(zip(SECTION_ORDER, data["section_scores"]), key=lambda item: item[1])
    weakest = min(zip(SECTION_ORDER, data["section_scores"]), key=lambda item: item[1])
    comparisons = benchmark_summary.get("comparisons", [])
    benchmark_line = "Benchmark comparison was limited by the data supplied."
    for item in comparisons:
        if item["status"] != "In line":
            benchmark_line = f"{item['label']} is {item['comment'].lower()}."
            break

    first_problem = report["top_problems"][0] if report["top_problems"] else "Control is uneven across the hiring process."
    first_action = report["day_30_plan"][0] if report["day_30_plan"] else "Tighten the weakest part of the process first."
    return [
        f"The strongest area is {strongest[0]} at {strongest[1]}/10.",
        f"The weakest area is {weakest[0]} at {weakest[1]}/10.",
        benchmark_line,
        f"First priority: {first_action.rstrip('.')}.",
    ]


def _section_rating(score: int) -> str:
    if score >= 8:
        return "Strong"
    if score >= 6:
        return "Sound with gaps"
    if score >= 4:
        return "Inconsistent"
    return "Weak"


def _status_fill(status: str) -> str:
    lowered = status.lower()
    if "ahead" in lowered:
        return "E7E7E7"
    if "behind" in lowered:
        return "F1F1F1"
    return "F7F7F7"


def _format_metric_value(value: float | None, suffix: str) -> str:
    if value is None:
        return "n/a"
    if float(value).is_integer():
        number = str(int(value))
    else:
        number = f"{value:.1f}".rstrip("0").rstrip(".")
    if suffix == "days":
        return f"{number} days"
    if suffix == "%":
        return f"{number}%"
    return number
